import os
import sys
from docx import Document, shared
from docx.shared import Inches, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import datetime
import re

class Templetes:
    def __init__(self):
        # self.wordDoc = Document(r'\Requirements.docx')
        # style = self.wordDoc.styles['Normal']
        # font = style.font
        # font.name = 'Times New Roman'
        pass


    def appendix_two(self, nom_akt='ГГГГ.ММ.ДД-ХХХХ ТБ'):
        self.wordDoc = Document(r'\Example.docx')
        style = self.wordDoc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = shared.Pt(11)

        p1 = self.wordDoc.paragraphs[0]
        p1.add_run('Text').italic = True
        p1.alignment = 2
        p2 = self.wordDoc.add_paragraph()
        p2.add_run('№' + nom_akt).italic = True
        p2.alignment = 2

        p3 = self.wordDoc.add_paragraph()
        p3.add_run('Дополнение TEXT2').bold = True
        p3.alignment = 1
        p3.italic = False

        p4 = self.wordDoc.add_paragraph()
        p4.add_run('Выводы о TEXT3').bold = True
        p4.alignment = 1
        p4.italic = False

        # p5 = self.wordDoc.add_paragraph()
        # p5.alignment = 0

    def tan(self):
        wd = self.wordDoc
        tabl_line = wd.add_table(2, 3)
        tabl_line.cell(0, 0).text = '____________________________'
        tabl_line.cell(0, 1).text = '____________________'
        tabl_line.cell(0, 2).text = '___________________________'
        tabl_line.cell(1, 0).text = 'наименование должности'
        tabl_line.cell(1, 1).text = 'подпись'
        tabl_line.cell(1, 2).text = 'расшифровка подписи'
        for i in range(3):
            tabl_line.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            tabl_line.cell(0, i).paragraphs[0].alignment = 1
            tabl_line.cell(0, i).paragraphs[0].runs[0].font.underline = True
            tabl_line.cell(0, i).paragraphs[0].paragraph_format.space_after = Pt(0)
            tabl_line.cell(1, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            tabl_line.cell(1, i).paragraphs[0].paragraph_format.space_after = Pt(0)
            tabl_line.cell(1, i).paragraphs[0].alignment = 1

    def page_analiz_skad(self, nom='YYYY.mm.dd-####-№№', path='', name='Анализ - СКАД', head=None):
        self.wordDoc = Document(r'Analiz_SKAD.docx')
        style = self.wordDoc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        wd = self.wordDoc
        """Шаблон формирования шапки файла СКАД"""

        self.wordDoc.sections[0].header.is_linked_to_previous = False

        tx_all = self.wordDoc.sections[0].header.paragraphs[0].text
        tx_one = self.wordDoc.sections[0].first_page_header.paragraphs[0].text

        self.wordDoc.sections[0].header.paragraphs[0].text = tx_all.replace('№YYYY.mm.dd-№№№№-## ', nom)
        self.wordDoc.sections[0].header.paragraphs[0].runs[0].font.size = Pt(10)
        self.wordDoc.sections[0].first_page_header.paragraphs[0].text = tx_one.replace('№YYYY.mm.dd-№№№№-## ', nom)
        self.wordDoc.sections[0].header.paragraphs[0].paragraphs[0].runs[0].font.size = Pt(10)

        tab = wd.add_table(2, 1)
        tab.cell(0, 0).text = head[0][0]
        tab.cell(0, 0).paragraphs[0].alignment = 2
        tab.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        tab.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)
        tab.cell(0, 1).text = head[1][0]
        tab.cell(0, 1).paragraphs[0].alignment = 2
        tab.cell(0, 1).paragraphs[0].runs[0].font.bold = True

        t1 = [[
                  '1. Проведенные мероприятия по выявлению новых видов и случаев противоправных действий со сотрудников в отношении клиентов:'],
              ['Анализ выставленных отклонений в СКАД ']]

        tab2 = wd.add_table(2, 1)
        tab2.style = 'TableGrid'
        tab2.cell(0, 0).text = t1[0][0]
        tab2.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        tab2.cell(0, 0).paragraphs[0].paragraph_format.first_line_indent = Inches(0.25)
        shadding_elm = parse_xml(r'<w:shd {} w:fill="A9A9A9"/>'.format(nsdecls('w')))
        tab2.cell(0, 0)._tc.get_or_add_tcPr().append(shadding_elm)
        tab2.cell(0, 1).text = t1[1][0]
        tab2.cell(0, 1).paragraphs[0].paragraph_format.first_line_indent = Inches(0.25)
        shadding_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        tab2.cell(0, 1)._tc.get_or_add_tcPr().append(shadding_elm)

    def skad_body(self, data):
        """Шаблон заполение основной части файла СКАД"""
        try:
            column = len(data)
            row = max(len(elem) for elem in data)
            doc = self.wordDoc

            table = doc.tables[2]
            print(table.cell(0, 0).paragraphs[0].text)
            tab1 = table.cell(0, 1).add_table(column, row)
            tab1.style = 'TableGrid'

            for i in range(len(data)):
                for j in range(len(data[i])):
                    tab1.cell(i, j).text = data[i][j]

            for i in range(len(data[0])):
                tab1.cell(0, i).paragraphs[0].runs[0].font.bold = True

            self.delete_paragraph(table.cell(0, 1).paragraphs[0])
            for row in tab1.rows:
                for cell in row.cells:
                    parags = cell.paragraphs
                    for paragraph in parags:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(8)
        except TypeError:
            print("Неверные входные значения")

    def check_list_doc(self):
        self.wordDoc = Document(r'\Check_list.docx')
        style = self.wordDoc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        """Шапка чек листа"""
        h1 = [['', 'ЧЕК-ЛИСТ', '']]
        self.creatTable(h1, firstlineindent=0)
        self.tab.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)
        self.tab.cell(0, 1).paragraphs[0].paragraph_format.space_after = Pt(0)
        self.tab.cell(0, 2).paragraphs[0].paragraph_format.space_after = Pt(0)
        self.paragraphs_alignment([[0, 1]], 1)
        self.cell_bold([[0, 1]], True)
        self.delete_paragraph(self.wordDoc.paragraphs[0])
        self.delete_paragraph(self.wordDoc.paragraphs[0])

    def creatTable(self, data=None, name_table="", autoformat=0, firstlineindent=0.25, check=False):
        try:
            p6 = self.wordDoc.add_paragraph()
            self.column = len(data)
            self.row = max(len(elem) for elem in data)
            p6.add_run(name_table)
            p6.paragraph_format.first_line_indent = Inches(firstlineindent)
            p6.alignment = 3
            self.tab = self.wordDoc.add_table(self.column, self.row)
            # при autoformat = 1 стиль таблицы с видимыми границами, иначе без границ
            if autoformat == 0:
                self.tab.style = 'WithoutBorders'
            elif autoformat == 1:
                self.tab.style = 'TableGrid'
            elif autoformat == 2:
                self.tab.style = 'OuterMarginOnly'

            for i in range(len(data)):
                for j in range(len(data[i])):
                    cell = self.tab.cell(i, j)

                    run = cell.paragraphs[0]  # add_paragraph()
                    run.paragraph_format.first_line_indent = Inches(firstlineindent)
                    run.add_run(data[i][j])
                    # print(i, j, data[i][j])

                    if check & (j > 1):
                        self.paragraphs_alignment([[i, j]], 1)
                        # self.cell_vertical_alignment([i + 1, j + 1])
                        self.tab.rows[i].cells[j].width = Inches(0.9)
                        if data[i][j] == '+':
                            self.color_cells([[i, j]], color="F08080")

                    if data[i][j] == 'Проведен анализ обращений по каналу ВСП:':
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].paragraph_format.first_line_indent = Inches(0)
        except TypeError:
            print("Неверные входные значения")

    def save_doc(self, path_file, file_name):
        """Модуль сохраняет документ. Если документ с таким именем есть, добавляет индекс к названию документа."""
        file_name = '\\' + file_name
        try:
            #  Добавленно при создлание актов, создает папку если такой не существует(исключение FileNotFoundError неотработает)
            if not os.path.exists(path_file):
                os.makedirs(path_file)
            #  ___________________________________
            if os.path.isfile(path_file + file_name + '.docx'):
                i = 1
                while os.path.isfile(path_file + file_name + '(' + str(i) + ')' + '.docx'):
                    i += 1
                new_path = path_file + file_name + '(' + str(i) + ')' + '.docx'
            else:
                new_path = path_file + file_name + '.docx'

            # Строчка отвечающая за то, куда сохраняется файл
            self.wordDoc.save(new_path)
            print("Файл сохранен", new_path)
            return new_path

        except FileNotFoundError:
            sys.exit("Папки не существует")

    def delete_paragraph(self, paragraph):
        """Модуль для удаления параграфа"""
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # сделать ячейку жирным шрифтом
    def cell_bold(self, cells=None, arg=False):
        """Модуль для выделения текста в ячейки жирным шфиртом.
        В качестве аргументов список списков из пары чисел(ячейки, индексирования с 0)
        и arg(True - жирный, False - не жирный)
        Пример cell_bold(cells=[[столбец, строка],[столбец, строка]], arg=True)"""
        try:
            for ind in cells:
                self.tab.cell(ind[0], ind[1]).paragraphs[0].runs[0].font.bold = arg
        except:
            print('Ошибка при выделении жирным шрифтом')
            pass

    # сделать ячейку наклонным шрифтом
    def cell_italic(self, cells=None, arg=False):
        """Модуль для выделения текста в ячейки наклонным шфиртом.
        В качестве аргументов список списков из пары чисел(ячейки[[столбец, строка],[столбец, строка]], индексирования с 0)
        и arg(True - наклонный, False - не наклонный)
        Пример cell_italic(cells=[[столбец, строка],[столбец, строка]], arg=True)"""
        try:
            for ind in cells:
                self.tab.cell(ind[0], ind[1]).paragraphs[0].runs[0].font.italic = arg
        except:
            print('Ошибка при выделении жирным шрифтом')
            pass

    # 0-net 1-da
    def cell_font_underline(self, cells=None, arg=False):
        """Модуль для подчеркивания текста в ячейки(-ах). В качестве аргументов список из последовательность пары чисел и arg где 0 - не подчеркивать, 1 - подчеркивать. Пример [1,1,1,2], подчеркнуться записи в ячейках (1,1) и (1,2).
        Пример cell_font_underline(cell=[2,3], arg=1), текст в ячейке (2,3) будет подчеркнут.
        Пример cell_font_underline(cell=[2,3, 4,2], arg=1), текст в ячейках (2,3) и (4,2) будут подчеркнуты
        Пример cell_font_underline(cell=[4,2], arg=0), текст в ячейке (4,2) будет без подчеркивания"""
        try:
            for ind in cells:
                self.tab.cell(ind[0], ind[1]).paragraphs[0].runs[0].font.underline = arg
        except:
            print('Ошибка при выделении жирным шрифтом')
            pass

    # позиционирование текста 1-центр 2-право
    def paragraphs_alignment(self, cells=None, arg=1):
        """Модуль для установки горизонтального позиционирования текста в ячейках. В качестве аргументов список списков из пары чисел и arg(позиционирование текста по горизонтали # 0-лево 1-центр 3-право).
        Пример paragraphs_alignment(cell=[[1,1],[1,2],[1,3]], arg=1), позиционирование текста в ячейках (1,1), (1,2), (1,3) по горизонтали будет по центру"""
        try:
            for ind in cells:
                self.tab.cell(ind[0], ind[1]).paragraphs[0].alignment = arg
        except:
            print('Ошибка при позиционировании по горизонтали')
            pass

    def color_cells(self, cells=None, color="A9A9A9"):
        """Модуль для изменения цвета ячеек или ячейки. В качестве аргументов список списков из пары чисел и цвет. Пример color_cells(cell=[[1,1],[1,2]], color="D6D6D6"), ячейки (1,1) и (1,2) будут покрашены"""
        try:
            for ind in cells:
                shadding_elm = parse_xml(r'<w:shd {0} w:fill="{1}"/>'.format(nsdecls('w'), color))
                self.tab.cell(ind[0], ind[1])._tc.get_or_add_tcPr().append(shadding_elm)
        except:
            print('Ошибка при окрашивание ячейки')
            pass

    def change_width(self, column, inches):
        for row in self.tab.rows:
            row.cells[column].width = Inches(inches)

if __name__ == '__main__':
    info = ['Галошин Александр Максимович', 'SD123456', 'ВСП', 'Нет', 'с 20.12.19 по 24.12.19']

    # body = [['Наименование продукта', 'Наименование алгоритмов', 'СККО', 'ККАФ'],
    #         ['Кредиты', 'Алгоритм 4.1',	'+', ''],
    #         ['Кредиты', 'Алгоритм 4.2',	'+', '+'],
    #         ['Кредиты', 'Алгоритм 4.3',	'+', ''],
    #         ['Дебетовые карты', 'Алгоритм 2.1',	'+', '+'],
    #         ['Дебетовые карты', 'Алгоритм 2.2',	'+', ''],
    #         ]

    test_all = [['Наименование продукта', 'Наименование алгоритмов', 'СККО', 'ККАФ'],
                ['Вклады', 'Алгоритм 1.1', '+', ''],
                ['Вклады', 'Алгоритм 1.2', '+', '+'],
                ['Вклады', 'Алгоритм 1.3', '+', ''],
                ['Вклады', 'Алгоритм 1.4', '', ''],
                ['Вклады', 'Алгоритм 1.5', '+', ''],
                ['Вклады', 'Алгоритм 1.6', '', ''],
                ['Вклады', 'Алгоритм 1.7', '', ''],
                ['Вклады', 'Алгоритм 1.8', '+', ''],
                ['Вклады', 'Алгоритм 1.9', '+', ''],
                ['Вклады', 'Алгоритм 1.10', '+', '+'],
                ['Вклады', 'Алгоритм 1.11', '', ''],
                ['Вклады', 'Алгоритм 1.12', '', ''],
                ['Вклады', 'Алгоритм 1.13', '+', '+'],
                ['Вклады', 'Алгоритм 1.14', '', '+'],
                ['Вклады', 'Алгоритм 1.15', '', ''],
                ['Вклады', 'Алгоритм 1.16', '', '+'],
                ['Вклады', 'Алгоритм 1.17', '+', '+'],
                ['Вклады', 'Алгоритм 1.18', '', ''],
                ['Вклады', 'Алгоритм 1.19', '+', ''],
                ['Вклады', 'Алгоритм 1.20', '+', '+'],
                ['Кредитные карты', 'Алгоритм 2.1', '+', '+'],
                ['Кредитные карты', 'Алгоритм 2.2', '', ''],
                ['Кредитные карты', 'Алгоритм 2.3', '+', ''],
                ['Кредитные карты', 'Алгоритм 2.4', '', ''],
                ['Кредитные карты', 'Алгоритм 2.5', '', ''],
                ['Кредитные карты', 'Алгоритм 2.6', '+', '+'],
                ['Кредитные карты', 'Алгоритм 2.7', '', '+'],
                ['Кредитные карты', 'Алгоритм 2.8', '+', ''],
                ['Кредитные карты', 'Алгоритм 2.9', '', ''],
                ['Кредитные карты', 'Алгоритм 2.10', '', ''],
                ['Дебетовые карты', 'Алгоритм 3.1', '+', ''],
                ['Дебетовые карты', 'Алгоритм 3.2', '', '+'],
                ['Дебетовые карты', 'Алгоритм 3.3', '', ''],
                ['Дебетовые карты', 'Алгоритм 3.4', '', ''],
                ['Дебетовые карты', 'Алгоритм 3.5', '+', '+'],
                ['Дебетовые карты', 'Алгоритм 3.6', '', ''],
                ['Дебетовые карты', 'Алгоритм 3.7', '', ''],
                ['Дебетовые карты', 'Алгоритм 3.8', '+', '+'],
                ['Дебетовые карты', 'Алгоритм 3.9', '', '+'],
                ['Дебетовые карты', 'Алгоритм 3.10', '+', ''],
                ['Кредиты', 'Алгоритм 4.1', '+', '+'],
                ['Кредиты', 'Алгоритм 4.1', '', ''],
                ['Кредиты', 'Алгоритм 4.1', '', ''],
                ['Кредиты', 'Алгоритм 4.1', '', '+'],
                ['Кредиты', 'Алгоритм 4.1', '+', ''],
                ['Кредиты', 'Алгоритм 4.1', '', ''],
                ['Кредиты', 'Алгоритм 4.1', '', ''],
                ['Кредиты', 'Алгоритм 4.1', '+', '+'],
                ]

    period_s = '2017-12-18'
    period_do = '2017-12-30'
    fio = ' Галошин Валерия Игоревна'
    tabel = '2113131'

    info_r = [
        ['ФИО сотрудника', fio],
        ['Табельный номер', tabel],
        ['Период', 'с ' + period_s + ' по ' + period_do]
    ]


    path_check_list = 'C:\\Users\\Galoshin\\i\\'  # папка для поиска актов возражений

    word_tabl = Templetes()
    word_tabl.check_list_doc()  # шапка файла
    word_tabl.creatTable(info_r, firstlineindent=0)
    word_tabl.change_width(0, 1.75)
    word_tabl.creatTable(test_all, autoformat=1, firstlineindent=0, check=True)  # значение таблицы(список списков)
    word_tabl.change_width(0, 2.2)
    word_tabl.change_width(1, 3)

    word_tabl.creatTable([['Целесообразность проведения: ' + 'Да/Нет']], autoformat=0)
    word_tabl.save_doc(path_check_list, 'Чек лист по - ' + fio) # путь и имя сохраняемого файла
